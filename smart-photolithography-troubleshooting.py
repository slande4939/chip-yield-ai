#!/usr/bin/env python
# coding: utf-8

# In[2]:


# Re-import necessary libraries since execution state was reset
import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Data for the Defect Type Breakdown Table
data = {
    "Defect Type": ["CD Variation", "Alignment Drift", "Particle Contamination"],
    "Root Cause": [
        "Resist thickness variation",
        "Stepper calibration issues",
        "Developer instability"
    ],
    "Fix Implemented": [
        "Optimized spin speed & bake temp",
        "Overlay correction & periodic reticle calibration",
        "Improved filtration & dispense control"
    ],
    "Impact on Yield": [
        "Yield improved by 10%",
        "15% overlay accuracy increase",
        "Defect density reduced by 30%"
    ]
}

# Create DataFrame
df = pd.DataFrame(data)

# Create a Word document
doc = Document()
doc.add_heading('Photolithography Defect Type Breakdown Table', level=1)

# Add table with gridlines
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'  # Apply table grid style for gridlines
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df.columns):
    hdr_cells[i].text = column_name

# Add data rows
for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Manually apply border style if needed for custom styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/></w:tcBorders>'.format(nsdecls('w'))))

# Save the document
doc_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(doc_path)

doc_path


# In[9]:


# Re-import necessary libraries since execution state was reset
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create a Word document
doc = Document()
doc.add_heading('Case Study 1: Overlay Misalignment Issue at Intel', level=1)

# Define data for the table
data = [
    ("Troubleshooting Breakdown", "Details"),
    ("Problem", "Yield loss due to overlay misalignment affecting pattern fidelity"),
    ("Root Cause Analysis", "SPC data indicated thermal expansion in reticle stage, "
                            "misalignment drift measured via overlay metrology"),
    ("Fix Implemented", "Implemented real-time stage compensation + periodic recalibration"),
    ("Results", "20% improvement in overlay accuracy, reduced rework costs"),
    ("Team Collaboration", "Worked with process integration & metrology teams to validate corrections")
]

# Add table to the document
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = data[0][0]
hdr_cells[1].text = data[0][1]

# Add data rows
for row_data in data[1:]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# Apply border styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}><w:top w:val="single" w:sz="4"/>'
            r'<w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/>'
            r'<w:right w:val="single" w:sz="4"/></w:tcBorders>'.format(nsdecls('w'))))

# Save the document
doc_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(doc_path)

doc_path


# In[1]:


from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create a Word document
doc = Document()
doc.add_heading('Case Study 2: Photoresist Coating Defects – Non-Uniformity & Contamination', level=1)

# Add table with gridlines
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'

# Add table headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Troubleshooting Breakdown"
hdr_cells[1].text = "Details"

# Add Problem row
problem_cells = table.rows[1].cells
problem_cells[0].text = "Problem"
problem_cells[1].text = "Yield loss due to non-uniform resist thickness and contamination spots, causing CD variation in critical layers."

# Add Root Cause Analysis row
root_cause_cells = table.rows[2].cells
root_cause_cells[0].text = "Root Cause Analysis"
root_cause_cells[1].text = (
    " SPC Data Analysis: Detected variation in coat thickness.\n"
    " CDSEM & Metrology: Identified micro-bubble defects.\n"
    " Process Logs: Found inconsistent resist dispense pressure."
)

# Apply border styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}>'
            r'<w:top w:val="single" w:sz="4"/>'
            r'<w:left w:val="single" w:sz="4"/>'
            r'<w:bottom w:val="single" w:sz="4"/>'
            r'<w:right w:val="single" w:sz="4"/>'
            r'</w:tcBorders>'.format(nsdecls('w'))
        ))

# Save the document
doc_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(doc_path)

doc_path


# In[2]:


# Re-import necessary libraries after execution state reset
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Create a Word document
doc = Document()
doc.add_heading('Case Study 2 – Fix Implemented & Results', level=1)

# Add table with two columns
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Define headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Fix Implemented"
hdr_cells[1].text = "Results"

# Define data
data = [
    ("Resist Viscosity Optimization", "Adjusted dispense pressure & flow rate."),
    ("Coat Track Tuning", "Optimized spin profile & dispense timing."),
    ("Filtration Upgrade", "Added a 0.1-micron resist filter."),
    ("Resist Uniformity", "Improved by 25%."),
    ("Yield", "Increased by 8% in critical patterned layers."),
    ("SPC Control Limits", "Stabilized, reducing process excursions.")
]

# Populate table
for fix, result in data:
    row_cells = table.add_row().cells
    row_cells[0].text = fix
    row_cells[1].text = result

# Apply border styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}>'
                                                         r'<w:top w:val="single" w:sz="4"/>'
                                                         r'<w:left w:val="single" w:sz="4"/>'
                                                         r'<w:bottom w:val="single" w:sz="4"/>'
                                                         r'<w:right w:val="single" w:sz="4"/>'
                                                         r'</w:tcBorders>'.format(nsdecls('w'))))

# Save document
file_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(file_path)

file_path


# In[1]:


# Re-import necessary libraries since the execution state was reset
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create a Word document
doc = Document()
doc.add_heading('Case Study 3 – Problem & Root Cause Analysis', level=1)

# Add table with gridlines
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Define headers
headers = ["Troubleshooting Breakdown", "Details"]
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Define data for the table
data = [
    ("Problem", "Overlay misalignment was detected in DUV stepper process, leading to pattern shift between layers.\n"
                "Affected high-density interconnect layers, causing 20% of wafers to fail electrical test."),
    ("Root Cause Analysis", " SPC Data Trend Analysis: Identified progressive overlay drift in stepper tool.\n"
                            " Reticle Inspection: Found thermal expansion effects on reticle stage, shifting alignment over time.\n"
                            " Tool Performance Logs: Detected wafer chuck vacuum degradation, impacting wafer positioning repeatability.")
]

# Add data rows
for row_data in data:
    row_cells = table.add_row().cells
    for i, value in enumerate(row_data):
        row_cells[i].text = value

# Apply border styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}><w:top w:val="single" w:sz="4"/>'
            r'<w:left w:val="single" w:sz="4"/>'
            r'<w:bottom w:val="single" w:sz="4"/>'
            r'<w:right w:val="single" w:sz="4"/></w:tcBorders>'.format(nsdecls('w'))
        ))

# Save the document
doc_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(doc_path)

# Provide the download link
doc_path


# In[6]:


from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Create a Word document
doc = Document()
doc.add_heading('Case Study 3 – Fix Implemented & Results', level=1)

# Add table
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'  # Apply table grid style for gridlines

# Define the table headers and content
content = [
    ("Fix Implemented", """\
    🔹 Stepper Calibration Update: Implemented real-time overlay correction model based on wafer stage feedback.
    🔹 Wafer Chuck Maintenance: Replaced vacuum system components to improve wafer hold stability.
    🔹 Reticle Stage Compensation Algorithm: Introduced a thermal drift compensation model to adjust alignment dynamically."""),
    
    ("Results", """\
    ✔ Overlay accuracy improved by 18%, reducing pattern shift defects.
    ✔ Electrical test pass rate increased by 12%.
    ✔ Reduced rework & scrap costs by $200K per quarter."""),

    ("Team Involvement", """\
    🔹 Collaborated with stepper tool engineers to update calibration settings.
    🔹 Worked with metrology team to implement overlay correction strategies.
    🔹 Coordinated with yield enhancement engineers to track electrical test improvements."""),
]

# Populate table
for row_idx, (header, details) in enumerate(content):
    table.cell(row_idx, 0).text = header
    table.cell(row_idx, 1).text = details

# Apply border styling
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}><w:top w:val="single" w:sz="4"/>'
            r'<w:left w:val="single" w:sz="4"/>'
            r'<w:bottom w:val="single" w:sz="4"/>'
            r'<w:right w:val="single" w:sz="4"/></w:tcBorders>'.format(nsdecls('w'))
        ))

# Save the document
file_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx" 
doc.save(file_path)

file_path


# In[3]:


import matplotlib.pyplot as plt

# Define categories and causes
categories = {
    "Equipment": ["Thermal expansion in reticle stage", "Wafer chuck misalignment"],
    "Process": ["Improper alignment calibration", "Inconsistent overlay correction"],
    "Material": ["Defective materials used in lithography"],
    "Environment": ["Wafer warpage due to temperature fluctuations"]
}

# Set up figure and axis
fig, ax = plt.subplots(figsize=(10, 6))
ax.set_xlim(-1, 6)
ax.set_ylim(-2, 2)
ax.axis("off")

# Draw the main spine
ax.plot([-0.5, 5.5], [0, 0], "k-", linewidth=2)

# Draw category branches
y_positions = [1.5, 0.8, -0.8, -1.5]  # Position of branches
x_start = 2.5  # Point where branches start

for i, (category, causes) in enumerate(categories.items()):
    y = y_positions[i]
    
    # Draw category line
    ax.plot([x_start, x_start + 1], [0, y], "k-", linewidth=2)
    
    # Add category label
    ax.text(x_start + 1.2, y, category, fontsize=12, fontweight="bold", verticalalignment="center")
    
    # Add causes
    for j, cause in enumerate(causes):
        ax.text(x_start + 2, y - j * 0.3, f"- {cause}", fontsize=10, verticalalignment="center")

# Add the problem statement at the center
ax.text(-0.4, 0, "Overlay Misalignment Issue", fontsize=14, fontweight="bold", verticalalignment="center")

plt.title("Fishbone Diagram for Overlay Misalignment", fontsize=14, fontweight="bold")
plt.show()


# In[2]:


import matplotlib.pyplot as plt

# Define categories and causes for AI-Driven Defect Reduction
categories = {
    "Equipment": ["Sensor miscalibration in inspection tools", "Hardware degradation over time"],
    "Process": ["Incorrect process parameters", "Variability in defect classification thresholds"],
    "Material": ["Impurities in raw materials", "Variations in chemical composition"],
    "Environment": ["Ambient temperature fluctuations", "Electrostatic discharge (ESD) interference"]
}

# Set up figure and axis
fig, ax = plt.subplots(figsize=(10, 6))
ax.set_xlim(-1, 6)
ax.set_ylim(-2, 2)
ax.axis("off")

# Draw the main spine
ax.plot([-0.5, 5.5], [0, 0], "k-", linewidth=2)

# Draw category branches
y_positions = [1.5, 0.8, -0.8, -1.5]  # Position of branches
x_start = 2.5  # Point where branches start

for i, (category, causes) in enumerate(categories.items()):
    y = y_positions[i]
    
    # Draw category line
    ax.plot([x_start, x_start + 1], [0, y], "k-", linewidth=2)
    
    # Add category label
    ax.text(x_start + 1.2, y, category, fontsize=12, fontweight="bold", verticalalignment="center")
    
    # Add causes
    for j, cause in enumerate(causes):
        ax.text(x_start + 2, y - j * 0.3, f"- {cause}", fontsize=10, verticalalignment="center")

# Add the problem statement at the center
ax.text(-0.4, 0, "AI-Driven Defect Reduction Issue", fontsize=14, fontweight="bold", verticalalignment="center")

plt.title("Fishbone Diagram for AI-Driven Defect Reduction", fontsize=14, fontweight="bold")
plt.show()


# In[15]:


import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Define data for the AI-Driven Defect Reduction case study
data = {
    "Troubleshooting Breakdown": [
        "Problem",
        "Root Cause Analysis",
        "Fix Implemented",
        "Results",
        "Team Collaboration"
    ],
    "Details": [
        "High defect rates impacting semiconductor yield and reliability.",
        "AI-driven analysis identified pattern defects, misalignment, and contamination as primary sources of yield loss.",
        "Applied AI classification models for early defect detection and integrated real-time correction mechanisms.",
        "25% reduction in defect rates, significant cost savings in rework and scrap.",
        "Worked with AI engineers, process integration, and metrology teams to optimize detection and prevention strategies."
    ]
}

# Create a DataFrame
df = pd.DataFrame(data)

# Create a Word document
doc = Document()
doc.add_heading('Case Study 2: AI-Driven Defect Reduction', level=1)
doc.add_paragraph("Table 2: Photolithography AI-Driven Defect Reduction Troubleshooting Summary")

# Add table with gridlines
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'  # Apply table grid style for gridlines

# Add headers
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df.columns):
    hdr_cells[i].text = column_name

# Add data rows
for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Apply border styles to each cell
for row in table.rows:
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(parse_xml(
            r'<w:tcBorders {}><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/></w:tcBorders>'.format(nsdecls('w'))
        ))

# Save the document
doc_path = "C:/Users/Lavieestbelle$1/Desktop/Biodegradable_Plastics_Project/Biodegradable_Plastics_Research_Table_with_Grid.docx"
doc.save(doc_path)

print(f"Document saved as {doc_path}")


# In[ ]:




